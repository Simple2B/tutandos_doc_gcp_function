from docx import Document
from .data_class import DataItem
from .utils import to_percentage
from .style_map import Styles


class DocumentGenerator():
    def __init__(self, template_path: str, output_path: str, data: DataItem):
        self.template_path = template_path
        self.output_path = output_path
        self.data = data
        self.doc = Document(template_path)

    def go(self):
        self._do_exec_summary()
        self._do_summary_of_results()
        self._do_details()
        self._do_sections()

        self._save()

    def _do_exec_summary(self):
        self.doc.add_heading('Executive Summary', 1)

        p = self.doc.add_paragraph('This report and associated site audit data has been complied by ', style=Styles.summary)
        p.add_run(self.data.metadata.client_name).bold = True
        p.add_run(' to review the security provision across the ')
        p.add_run(self.data.metadata.location).bold = True
        p.add_run(' site in accordance with the ENGIE Site Risk Assessment Framework. The site audit was conducted on ')
        p.add_run(self.data.metadata.date).bold = True
        p.add_run('. The staff on site for the audit were first class and helped facilitate a straightforward and well informed audit.  ')
        self.doc.add_page_break()

    def _do_summary_of_results(self):
        self.doc.add_heading('Summary of Results', 1)
        self.doc.add_paragraph('Results here with table', style=Styles.summary)

        table = self.doc.add_table(rows=1, cols=3, style='Plain Table 1')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Section'
        hdr_cells[1].text = 'Security Dimension'
        hdr_cells[2].text = 'Score'
        for sectionData in self.data.sectionDataList:
            row_cells = table.add_row().cells
            row_cells[0].text = str(sectionData.number)
            row_cells[1].text = sectionData.title
            row_cells[2].text = to_percentage(sectionData.score)

        self.doc.add_page_break()

    def _do_details(self):
        self.doc.add_heading('Site and Audit Details', 1)
        self.doc.add_heading(self.data.metadata.location, 2)
        self.doc.add_heading(self.data.metadata.client_name + ' | ' + self.data.metadata.date, 3)

        records = (
            ('Conducted on', '', self.data.metadata.date),
            ('Client Name', '', self.data.metadata.client_name),
            ('Location Name', '', self.data.metadata.location)
        )

        table = self.doc.add_table(rows=1, cols=3, style='Plain Table 1')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Score'
        hdr_cells[1].text = 'Failed Items'
        hdr_cells[2].text = 'Actions'
        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc

        self.doc.add_page_break()

    def _subsection_filter(self, subsection, section_number):
        return subsection.number.startswith(section_number)

    def _do_sections(self):
        self.doc.add_heading('Site Security Survey', 1)

        for section in self.data.sectionDataList:
            # print('section.number type = ' + str(type(section.number)))
            # print('section.title type = ' + str(type(section.title)))
            self.doc.add_heading(section.number + ' ' + section.title, 2)

            sublist = filter(lambda x : self._subsection_filter(x, section.number + '.'), self.data.subsectionDataList)

            # print('NEW SUBLIST=============-')
            for subsection in sublist:
                print(subsection)
                self.doc.add_heading(subsection.title, 3)
                self.doc.add_paragraph(str(subsection.score), style=Styles.subsection_para)
                self.doc.add_paragraph(subsection.comments, style=Styles.subsection_para)

        self.doc.add_page_break()

    def _save(self):
        self.doc.save(self.output_path)
